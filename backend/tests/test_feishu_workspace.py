import asyncio
import io
import zipfile
from docx import Document

from app.adapters.feishu import FeishuApiError
from app.core.config import settings
from app.services.feishu_workspace import FeishuWorkspaceService


class FakeFeishuClient:
    def __init__(self) -> None:
        self.moves = []
        self.uploads = []

    async def list_folder_items(self, folder_token: str, *, page_size: int = 200, page_token: str | None = None) -> dict:
        if folder_token == "parent_folder":
            return {
                "code": 0,
                "data": {
                    "files": [
                        {"name": "AI生成", "token": "new_root", "type": "folder", "url": "https://feishu.test/drive/folder/new_root"}
                    ],
                    "has_more": False,
                },
            }
        if folder_token == "new_root":
            return {
                "code": 0,
                "data": {
                    "files": [
                        {"name": "Deep Research", "token": "deep_research", "type": "folder", "url": "https://feishu.test/drive/folder/deep_research"},
                        {"name": "分镜项目", "token": "storyboards", "type": "folder", "url": "https://feishu.test/drive/folder/storyboards"},
                    ],
                    "has_more": False,
                },
            }
        if folder_token == "old_root":
            return {
                "code": 0,
                "data": {
                    "files": [
                        {"name": "旧文档", "token": "doc_1", "type": "docx"},
                        {"name": "旧表格", "token": "sheet_1", "type": "sheet"},
                    ],
                    "has_more": False,
                },
            }
        return {"code": 0, "data": {"files": [], "has_more": False}}

    async def move_file(self, file_token: str, *, folder_token: str, file_type: str = "file") -> dict:
        self.moves.append((file_token, folder_token, file_type))
        return {"code": 0, "data": {}}

    async def create_folder(self, parent_token: str, name: str) -> dict:
        token = {
            "AI生成": "new_root",
            "Deep Research": "deep_research",
            "分镜项目": "storyboards",
        }.get(name, "created_folder")
        return {"code": 0, "data": {"token": token, "url": f"https://feishu.test/drive/folder/{token}"}}

    async def upload_file(self, folder_token: str, name: str, content: bytes) -> dict:
        self.uploads.append((folder_token, name, content))
        return {"code": 0, "data": {"file_token": "file_123"}}

    async def get_document_raw_content(self, document_id: str) -> dict:
        return {"code": 0, "data": {"content": [{"type": "text", "text": "hello"}], "document_id": document_id}}

    async def download_drive_file(self, file_token: str) -> tuple[str, bytes, str]:
        return ("notes.md", b"# hello", "text/markdown")


def test_workspace_service_migrates_old_root_into_new_ai_folder(monkeypatch):
    monkeypatch.setattr(settings, "feishu_workspace_parent_url", "https://feishu.test/drive/folder/parent_folder")
    monkeypatch.setattr(settings, "feishu_workspace_folder_name", "AI生成")
    monkeypatch.setattr(settings, "feishu_root_folder_token", "old_root")
    service = FeishuWorkspaceService(feishu=FakeFeishuClient())

    result = asyncio.run(service.ensure_default_workspace_folder())

    assert result["folder_token"] == "new_root"
    assert result["moved_items"] == 2
    assert result["folder_url"] == "https://feishu.test/drive/folder/new_root"


def test_workspace_service_ignores_deleted_legacy_root(monkeypatch):
    class DeletedRootFeishuClient(FakeFeishuClient):
        async def list_folder_items(self, folder_token: str, *, page_size: int = 200, page_token: str | None = None) -> dict:
            if folder_token == "old_root":
                raise FeishuApiError("Feishu HTTP error: 404, msg=file has been delete.")
            return await super().list_folder_items(folder_token, page_size=page_size, page_token=page_token)

    monkeypatch.setattr(settings, "feishu_workspace_parent_url", "https://feishu.test/drive/folder/parent_folder")
    monkeypatch.setattr(settings, "feishu_workspace_folder_name", "AI生成")
    monkeypatch.setattr(settings, "feishu_root_folder_token", "old_root")
    service = FeishuWorkspaceService(feishu=DeletedRootFeishuClient())

    result = asyncio.run(service.ensure_default_workspace_folder())

    assert result["folder_token"] == "new_root"
    assert result["moved_items"] == 0


def test_workspace_service_saves_markdown_doc(monkeypatch):
    monkeypatch.setattr(settings, "feishu_root_folder_token", "old_root")
    monkeypatch.setattr(settings, "feishu_workspace_parent_url", "https://feishu.test/drive/folder/parent_folder")
    monkeypatch.setattr(settings, "feishu_workspace_folder_name", "AI生成")
    monkeypatch.setattr(settings, "feishu_workspace_deep_research_folder_name", "Deep Research")
    fake = FakeFeishuClient()
    service = FeishuWorkspaceService(feishu=fake)

    result = asyncio.run(service.save_markdown_document(title="报告", markdown="# 报告\n内容 [来源](https://example.com)"))

    assert result.document_id == "file_123"
    assert result.folder_token == "deep_research"
    assert result.url.endswith("/file/file_123")
    assert fake.moves == [("doc_1", "new_root", "docx"), ("sheet_1", "new_root", "sheet")]
    assert fake.uploads[0][0] == "deep_research"
    assert fake.uploads[0][1].endswith(".docx")
    with zipfile.ZipFile(io.BytesIO(fake.uploads[0][2])) as archive:
        assert "word/document.xml" in archive.namelist()
        assert "word/numbering.xml" in archive.namelist()
        document_xml = archive.read("word/document.xml").decode("utf-8")
        rels_xml = archive.read("word/_rels/document.xml.rels").decode("utf-8")
        assert "报告" in document_xml
        assert "内容" in document_xml
        assert 'Target="https://example.com"' in rels_xml


def test_workspace_service_renders_real_links_and_lists(monkeypatch):
    monkeypatch.setattr(settings, "feishu_root_folder_token", "old_root")
    monkeypatch.setattr(settings, "feishu_workspace_parent_url", "https://feishu.test/drive/folder/parent_folder")
    monkeypatch.setattr(settings, "feishu_workspace_folder_name", "AI生成")
    monkeypatch.setattr(settings, "feishu_workspace_deep_research_folder_name", "Deep Research")
    fake = FakeFeishuClient()
    service = FeishuWorkspaceService(feishu=fake)
    markdown = "\n".join(
        [
            "# 研究摘要",
            "- 第一条 [来源一](https://example.com/a)",
            "- 第二条",
            "1. 编号一",
            "2. 编号二",
        ]
    )

    asyncio.run(service.save_markdown_document(title="结构测试", markdown=markdown))

    with zipfile.ZipFile(io.BytesIO(fake.uploads[0][2])) as archive:
        document_xml = archive.read("word/document.xml").decode("utf-8")
        rels_xml = archive.read("word/_rels/document.xml.rels").decode("utf-8")
        numbering_xml = archive.read("word/numbering.xml").decode("utf-8")
        assert "<w:hyperlink " in document_xml
        assert "第一条" in document_xml
        assert "编号二" in document_xml
        assert 'Target="https://example.com/a"' in rels_xml
        assert 'w:numFmt w:val="bullet"' in numbering_xml
        assert 'w:numFmt w:val="decimal"' in numbering_xml


def test_workspace_service_renders_markdown_tables(monkeypatch):
    monkeypatch.setattr(settings, "feishu_root_folder_token", "old_root")
    monkeypatch.setattr(settings, "feishu_workspace_parent_url", "https://feishu.test/drive/folder/parent_folder")
    monkeypatch.setattr(settings, "feishu_workspace_folder_name", "AI生成")
    monkeypatch.setattr(settings, "feishu_workspace_deep_research_folder_name", "Deep Research")
    fake = FakeFeishuClient()
    service = FeishuWorkspaceService(feishu=fake)
    markdown = "\n".join(
        [
            "| 模块 | 状态 |",
            "| --- | --- |",
            "| 公司时间线 | 完成 |",
            "| 融资历史 | 待补充 |",
        ]
    )

    asyncio.run(service.save_markdown_document(title="表格测试", markdown=markdown))

    with zipfile.ZipFile(io.BytesIO(fake.uploads[0][2])) as archive:
        document_xml = archive.read("word/document.xml").decode("utf-8")
        assert "<w:tbl>" in document_xml
        assert "公司时间线" in document_xml
        assert "待补充" in document_xml


def test_workspace_service_renders_large_markdown_as_docx(monkeypatch):
    monkeypatch.setattr(settings, "feishu_root_folder_token", "old_root")
    monkeypatch.setattr(settings, "feishu_workspace_parent_url", "https://feishu.test/drive/folder/parent_folder")
    monkeypatch.setattr(settings, "feishu_workspace_folder_name", "AI生成")
    monkeypatch.setattr(settings, "feishu_workspace_deep_research_folder_name", "Deep Research")
    fake = FakeFeishuClient()
    service = FeishuWorkspaceService(feishu=fake)
    markdown = "\n".join(f"- 第 {index} 行" for index in range(1, 70))

    result = asyncio.run(service.save_markdown_document(title="长报告", markdown=markdown))

    assert result.document_id == "file_123"
    with zipfile.ZipFile(io.BytesIO(fake.uploads[0][2])) as archive:
        document_xml = archive.read("word/document.xml").decode("utf-8")
        assert "第 1 行" in document_xml
        assert "第 69 行" in document_xml


def test_workspace_service_reads_doc_and_file_sources():
    service = FeishuWorkspaceService(feishu=FakeFeishuClient())

    doc = asyncio.run(service.read_reference("https://feishu.test/docx/doc_abc"))
    file_item = asyncio.run(service.read_reference("https://feishu.test/file/file_abc"))

    assert doc["type"] == "feishu_doc"
    assert doc["content_json"]["document_id"] == "doc_abc"
    assert "hello" in doc["text_content"]
    assert file_item["type"] == "feishu_file"
    assert "# hello" in file_item["text_content"]


def test_workspace_service_decodes_docx_file_text():
    doc = Document()
    doc.add_heading("广告 brief", level=1)
    doc.add_paragraph("车辆外观登场")
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "镜头"
    table.cell(0, 1).text = "内容"
    table.cell(1, 0).text = "1"
    table.cell(1, 1).text = "夜景穿梭"
    buffer = io.BytesIO()
    doc.save(buffer)

    text = FeishuWorkspaceService(feishu=FakeFeishuClient())._decode_file_content(
        buffer.getvalue(),
        mime_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        filename="brief.docx",
    )

    assert "广告 brief" in text
    assert "车辆外观登场" in text
    assert "镜头 | 内容" in text
    assert "1 | 夜景穿梭" in text


def test_workspace_service_uploads_to_default_workspace_when_target_folder_is_missing(monkeypatch):
    monkeypatch.setattr(settings, "feishu_workspace_parent_url", "https://feishu.test/drive/folder/parent_folder")
    monkeypatch.setattr(settings, "feishu_workspace_folder_name", "AI生成")
    fake = FakeFeishuClient()

    async def failing_upload(folder_token: str, name: str, content: bytes) -> dict:
        fake.uploads.append((folder_token, name, content))
        if folder_token == "deleted_folder":
            raise FeishuApiError("Feishu HTTP error: 400, msg=parent node not exist.")
        return {"code": 0, "data": {"file_token": "file_456"}}

    fake.upload_file = failing_upload
    service = FeishuWorkspaceService(feishu=fake)

    result = asyncio.run(service.save_markdown_document(title="回退文档", markdown="# 回退", folder_token="deleted_folder"))

    assert result.document_id == "file_456"
    assert result.folder_token == "deep_research"
    assert [item[0] for item in fake.uploads] == ["deleted_folder", "deep_research"]


def test_workspace_service_save_markdown_document_falls_back_to_deep_research_folder(monkeypatch):
    monkeypatch.setattr(settings, "feishu_root_folder_token", "old_root")
    monkeypatch.setattr(settings, "feishu_workspace_parent_url", "https://feishu.test/drive/folder/parent_folder")
    monkeypatch.setattr(settings, "feishu_workspace_folder_name", "AI生成")
    monkeypatch.setattr(settings, "feishu_workspace_deep_research_folder_name", "Deep Research")
    fake = FakeFeishuClient()

    async def failing_upload(folder_token: str, name: str, content: bytes) -> dict:
        fake.uploads.append((folder_token, name, content))
        if folder_token == "deleted_folder":
            raise FeishuApiError("Feishu HTTP error: 400, msg=parent node not exist.")
        return {"code": 0, "data": {"file_token": "file_789"}}

    fake.upload_file = failing_upload
    service = FeishuWorkspaceService(feishu=fake)

    result = asyncio.run(service.save_markdown_document(title="目标目录回退", markdown="# 回退", folder_token="deleted_folder"))

    assert result.document_id == "file_789"
    assert result.folder_token == "deep_research"
    assert [item[0] for item in fake.uploads] == ["deleted_folder", "deep_research"]
