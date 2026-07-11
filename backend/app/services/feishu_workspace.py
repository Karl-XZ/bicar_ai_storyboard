from __future__ import annotations

import io
import json
import re
from dataclasses import dataclass
from pathlib import Path
from urllib.parse import urlparse
from xml.sax.saxutils import escape

import markdown as markdown_lib
from bs4 import BeautifulSoup, NavigableString, Tag
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt

from app.adapters.feishu import FeishuClient, FeishuApiError
from app.core.config import settings


@dataclass(frozen=True)
class FeishuDocumentResult:
    document_id: str
    url: str
    folder_token: str


@dataclass(frozen=True)
class _DocxHyperlink:
    rel_id: str
    url: str


class FeishuWorkspaceService:
    def __init__(self, feishu: FeishuClient | None = None) -> None:
        self.feishu = feishu or FeishuClient()

    async def ensure_default_workspace_folder(self) -> dict[str, str | int]:
        parent_token = self._folder_token_from_url(getattr(settings, "feishu_workspace_parent_url", "")) or "root"
        try:
            target_folder = await self.ensure_named_folder(parent_token=parent_token, name=settings.feishu_workspace_folder_name)
            target_token = self._extract_token(target_folder)
            moved_items = 0
            source_token = settings.feishu_root_folder_token
            if source_token and source_token != target_token:
                try:
                    moved_items = await self.move_all_items(source_folder_token=source_token, target_folder_token=target_token)
                except FeishuApiError:
                    moved_items = 0
            return {
                "folder_token": target_token,
                "folder_url": self._extract_url(target_folder) or self._drive_folder_url(target_token),
                "moved_items": moved_items,
            }
        except FeishuApiError:
            current_token = settings.feishu_root_folder_token or "root"
            return {
                "folder_token": current_token,
                "folder_url": self._drive_folder_url(current_token),
                "moved_items": 0,
            }

    async def ensure_storyboard_workspace_folder(self) -> dict[str, str | int]:
        return await self.ensure_workspace_subfolder(settings.feishu_workspace_storyboard_folder_name)

    async def ensure_deep_research_workspace_folder(self) -> dict[str, str | int]:
        return await self.ensure_workspace_subfolder(settings.feishu_workspace_deep_research_folder_name)

    async def ensure_workspace_subfolder(self, name: str) -> dict[str, str | int]:
        root = await self.ensure_default_workspace_folder()
        parent_token = str(root.get("folder_token") or settings.feishu_root_folder_token or "root")
        folder = await self.ensure_named_folder(parent_token=parent_token, name=name)
        folder_token = self._extract_token(folder)
        return {
            "folder_token": folder_token,
            "folder_url": self._extract_url(folder) or self._drive_folder_url(folder_token),
            "moved_items": 0,
        }

    async def ensure_named_folder(self, *, parent_token: str, name: str) -> dict:
        page_token: str | None = None
        while True:
            response = await self.feishu.list_folder_items(parent_token, page_size=200, page_token=page_token)
            data = response.get("data", {})
            items = data.get("files") or data.get("items") or []
            for item in items:
                if str(item.get("name") or "").strip() == name and self._item_type(item) == "folder":
                    return {"code": 0, "data": {"token": self._item_token(item), "url": self._item_url(item)}}
            page_token = data.get("next_page_token") or data.get("page_token")
            if not page_token or not data.get("has_more"):
                break
        return await self.feishu.create_folder(parent_token, name)

    async def move_all_items(self, *, source_folder_token: str, target_folder_token: str) -> int:
        moved = 0
        page_token: str | None = None
        while True:
            response = await self.feishu.list_folder_items(source_folder_token, page_size=200, page_token=page_token)
            data = response.get("data", {})
            items = data.get("files") or data.get("items") or []
            for item in items:
                token = self._item_token(item)
                if not token:
                    continue
                file_type = self._item_type(item)
                if file_type == "folder" and token == target_folder_token:
                    continue
                try:
                    await self.feishu.move_file(token, folder_token=target_folder_token, file_type=file_type)
                    moved += 1
                except FeishuApiError:
                    continue
            page_token = data.get("next_page_token") or data.get("page_token")
            if not page_token or not data.get("has_more"):
                break
        return moved

    async def save_markdown_document(self, *, title: str, markdown: str, folder_token: str | None = None) -> FeishuDocumentResult:
        target_folder = str(folder_token).strip() if folder_token else ""
        deep_research_folder = await self.ensure_deep_research_workspace_folder()
        fallback_folder = str(deep_research_folder.get("folder_token") or settings.feishu_root_folder_token or "root")
        if not target_folder:
            target_folder = fallback_folder
        filename = f"{title}.docx"
        content = self._render_markdown_docx(title=title, markdown=markdown)
        try:
            upload = await self.feishu.upload_file(target_folder, filename, content)
            resolved_folder = target_folder
        except FeishuApiError as exc:
            if not self._is_missing_parent_folder_error(exc):
                raise
            if not fallback_folder or fallback_folder == target_folder:
                raise FeishuApiError(f"Feishu HTTP error: missing folder and no fallback available for {target_folder}")
            upload = await self.feishu.upload_file(fallback_folder, filename, content)
            resolved_folder = fallback_folder
        file_token = str((upload.get("data") or {}).get("file_token") or (upload.get("data") or {}).get("file", {}).get("file_token") or "")
        return FeishuDocumentResult(
            document_id=file_token or title,
            url=self._drive_file_url(file_token) if file_token else self._drive_folder_url(resolved_folder),
            folder_token=resolved_folder,
        )

    async def save_text_file(self, *, filename: str, text: str, folder_token: str | None = None) -> FeishuDocumentResult:
        upload, resolved_folder = await self.upload_file_with_fallback(
            target_folder=folder_token,
            name=filename,
            content=(text or "").encode("utf-8"),
        )
        file_token = str((upload.get("data") or {}).get("file_token") or (upload.get("data") or {}).get("file", {}).get("file_token") or "")
        return FeishuDocumentResult(
            document_id=file_token or filename,
            url=self._drive_file_url(file_token) if file_token else self._drive_folder_url(resolved_folder),
            folder_token=resolved_folder,
        )

    async def copy_local_docx_to_workspace(
        self,
        *,
        source_path: str | Path,
        title: str,
        folder_token: str | None = None,
    ) -> FeishuDocumentResult:
        path = Path(source_path).expanduser()
        if not path.exists() or not path.is_file():
            raise FileNotFoundError(f"模板文件不存在：{path}")
        if path.suffix.lower() != ".docx":
            raise ValueError(f"模板文件必须是 .docx：{path.name}")
        filename = self._docx_filename(title)
        target_folder = folder_token or await self.default_upload_folder_token()
        upload, resolved_folder = await self.upload_file_with_fallback(
            target_folder=target_folder,
            name=filename,
            content=path.read_bytes(),
        )
        file_token = str((upload.get("data") or {}).get("file_token") or (upload.get("data") or {}).get("file", {}).get("file_token") or "")
        return FeishuDocumentResult(
            document_id=file_token or filename,
            url=self._drive_file_url(file_token) if file_token else self._drive_folder_url(resolved_folder),
            folder_token=resolved_folder,
        )

    def folder_token_from_url(self, url: str | None) -> str | None:
        return self._folder_token_from_url(url)

    async def upload_file_with_fallback(self, *, target_folder: str | None, name: str, content: bytes) -> tuple[dict, str]:
        primary_folder = str(target_folder or settings.feishu_root_folder_token or "root")
        try:
            return await self.feishu.upload_file(primary_folder, name, content), primary_folder
        except FeishuApiError as exc:
            if not self._is_missing_parent_folder_error(exc):
                raise
        fallback_folder = await self.default_workspace_folder_token()
        if not fallback_folder or fallback_folder == primary_folder:
            raise FeishuApiError(f"Feishu HTTP error: missing folder and no fallback available for {primary_folder}")
        return await self.feishu.upload_file(fallback_folder, name, content), fallback_folder

    async def create_folder_with_fallback(self, *, parent_token: str | None, name: str) -> tuple[dict, str]:
        primary_parent = str(parent_token or settings.feishu_root_folder_token or "root")
        try:
            return await self.feishu.create_folder(primary_parent, name), primary_parent
        except FeishuApiError as exc:
            if not self._is_missing_parent_folder_error(exc):
                raise
        fallback_parent = await self.default_workspace_folder_token()
        if not fallback_parent or fallback_parent == primary_parent:
            raise FeishuApiError(f"Feishu HTTP error: missing folder and no fallback available for {primary_parent}")
        return await self.feishu.create_folder(fallback_parent, name), fallback_parent

    async def create_bitable_with_fallback(self, *, folder_token: str | None, name: str) -> tuple[dict, str]:
        primary_folder = str(folder_token or settings.feishu_root_folder_token or "root")
        try:
            return await self.feishu.create_bitable_app(name, folder_token=primary_folder), primary_folder
        except FeishuApiError as exc:
            if not self._is_missing_parent_folder_error(exc):
                raise
        fallback_folder = await self.default_workspace_folder_token()
        if not fallback_folder or fallback_folder == primary_folder:
            raise FeishuApiError(f"Feishu HTTP error: missing folder and no fallback available for {primary_folder}")
        return await self.feishu.create_bitable_app(name, folder_token=fallback_folder), fallback_folder

    async def default_workspace_folder_token(self) -> str:
        ensured = await self.ensure_default_workspace_folder()
        return str(ensured.get("folder_token") or settings.feishu_root_folder_token or "root")

    async def default_upload_folder_token(self) -> str:
        parent_token = self._folder_token_from_url(getattr(settings, "feishu_workspace_parent_url", "")) or settings.feishu_root_folder_token or "root"
        try:
            target_folder = await self.ensure_named_folder(parent_token=parent_token, name=settings.feishu_workspace_folder_name)
            return self._extract_token(target_folder) or settings.feishu_root_folder_token or "root"
        except FeishuApiError:
            return settings.feishu_root_folder_token or "root"

    async def read_reference(self, url_or_token: str) -> dict:
        value = str(url_or_token or "").strip()
        document_id = self._document_id_from_url(value)
        if document_id:
            raw = await self.feishu.get_document_raw_content(document_id)
            content_json = raw.get("data") or raw
            return {
                "type": "feishu_doc",
                "document_id": document_id,
                "url": self._doc_url(document_id),
                "content_json": content_json,
                "text_content": self._decode_document_content(content_json),
            }
        file_token = self._file_token_from_url(value) or self._feishu_token(value)
        if file_token:
            filename, content, mime_type = await self.feishu.download_drive_file(file_token)
            return {
                "type": "feishu_file",
                "file_token": file_token,
                "filename": filename,
                "mime_type": mime_type,
                "text_content": self._decode_file_content(content, mime_type=mime_type, filename=filename),
            }
        raise RuntimeError("无法识别飞书文档或文件链接")

    def _folder_token_from_url(self, url: str | None) -> str | None:
        if not url:
            return None
        parsed = urlparse(str(url))
        parts = [part for part in parsed.path.split("/") if part]
        try:
            index = parts.index("folder")
            return parts[index + 1] if len(parts) > index + 1 else None
        except ValueError:
            return None

    def _document_id_from_url(self, url: str | None) -> str | None:
        if not url:
            return None
        parsed = urlparse(str(url))
        parts = [part for part in parsed.path.split("/") if part]
        for marker in ("docx", "docs"):
            if marker in parts:
                index = parts.index(marker)
                if len(parts) > index + 1:
                    return parts[index + 1]
        return None

    def _file_token_from_url(self, url: str | None) -> str | None:
        if not url:
            return None
        parsed = urlparse(str(url))
        parts = [part for part in parsed.path.split("/") if part]
        if "file" in parts:
            index = parts.index("file")
            if len(parts) > index + 1:
                return parts[index + 1]
        return None

    def _feishu_token(self, value: str) -> str | None:
        if value.startswith("feishu://"):
            return value.replace("feishu://", "", 1).strip().strip("/")
        return None

    def _extract_token(self, response: dict) -> str:
        data = response.get("data", {})
        return str(data.get("token") or data.get("folder_token") or data.get("node", {}).get("token") or "")

    def _extract_url(self, response: dict) -> str | None:
        data = response.get("data", {})
        return data.get("url") or data.get("node", {}).get("url")

    def _docx_filename(self, title: str) -> str:
        value = re.sub(r"\s+", " ", str(title or "").strip())
        value = re.sub(r'[\\/:*?"<>|]+', "_", value).strip(" .")
        if value.lower().endswith(".docx"):
            value = value[:-5].strip(" .")
        if not value:
            raise ValueError("文件名不能为空")
        return f"{value[:120]}.docx"

    def _is_missing_parent_folder_error(self, exc: Exception) -> bool:
        message = str(exc).lower()
        return "parent node not exist" in message or ("parent node" in message and "not exist" in message)

    def _extract_document_id(self, response: dict) -> str:
        data = response.get("data", {})
        document = data.get("document") or {}
        return str(document.get("document_id") or data.get("document_id") or data.get("token") or "")

    def _item_token(self, item: dict) -> str:
        return str(item.get("token") or item.get("file_token") or item.get("node_token") or "")

    def _item_type(self, item: dict) -> str:
        raw = str(item.get("type") or item.get("mime_type") or item.get("file_type") or "").lower()
        if raw in {"folder", "explorer"}:
            return "folder"
        if raw in {"doc", "docx", "sheet", "mindnote", "bitable"}:
            return raw
        return "file"

    def _item_url(self, item: dict) -> str | None:
        return item.get("url") or item.get("link")

    def _doc_url(self, document_id: str) -> str:
        return f"{self._feishu_site_url()}/docx/{document_id}"

    def _drive_folder_url(self, folder_token: str) -> str:
        return f"{self._feishu_site_url()}/drive/folder/{folder_token}"

    def _drive_file_url(self, file_token: str) -> str:
        return f"{self._feishu_site_url()}/file/{file_token}"

    def _feishu_site_url(self) -> str:
        for configured_url in (settings.feishu_workspace_parent_url, settings.debug_paper_target_folder_url):
            netloc = urlparse(str(configured_url or "")).netloc
            if netloc:
                return f"https://{netloc}"
        domain = settings.feishu_base_url.replace("https://open.", "").replace("http://open.", "")
        return f"https://{domain}"

    def _decode_file_content(self, content: bytes, *, mime_type: str, filename: str) -> str:
        suffix = Path(filename).suffix.lower()
        if suffix == ".docx":
            try:
                document = Document(io.BytesIO(content))
                paragraphs = [paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip()]
                table_rows: list[str] = []
                for table in document.tables:
                    for row in table.rows:
                        cells = [cell.text.strip() for cell in row.cells]
                        if any(cells):
                            table_rows.append(" | ".join(cells))
                combined = paragraphs + table_rows
                if combined:
                    return "\n".join(combined)
            except Exception:
                pass
        if mime_type.startswith("text/") or suffix in {".md", ".markdown", ".txt", ".json", ".csv", ".tsv"}:
            try:
                return content.decode("utf-8")
            except UnicodeDecodeError:
                return content.decode("utf-8", errors="replace")
        if suffix == ".json":
            try:
                return json.dumps(json.loads(content.decode("utf-8")), ensure_ascii=False, indent=2)
            except Exception:
                pass
        return f"[binary file omitted] mime_type={mime_type} filename={filename} bytes={len(content)}"

    def _decode_document_content(self, payload: dict | str | list | None) -> str:
        if payload is None:
            return ""
        if isinstance(payload, str):
            stripped = payload.strip()
            if not stripped:
                return ""
            try:
                parsed = json.loads(stripped)
            except Exception:
                return stripped
            return self._decode_document_content(parsed)

        text_chunks: list[str] = []
        seen: set[str] = set()

        def append(value: str) -> None:
            normalized = re.sub(r"\s+", " ", str(value or "")).strip()
            if not normalized or normalized in seen:
                return
            seen.add(normalized)
            text_chunks.append(normalized)

        def walk(value) -> None:
            if value is None:
                return
            if isinstance(value, str):
                append(value)
                return
            if isinstance(value, list):
                for item in value:
                    walk(item)
                return
            if not isinstance(value, dict):
                return

            for key in ("raw_content", "markdown", "text", "title", "header", "body", "content"):
                nested = value.get(key)
                if isinstance(nested, str):
                    append(nested)
                elif isinstance(nested, (list, dict)):
                    walk(nested)

            elements = (
                value.get("elements"),
                value.get("blocks"),
                value.get("children"),
                value.get("paragraphs"),
                value.get("paragraph"),
                value.get("rows"),
                value.get("cells"),
            )
            for nested in elements:
                if nested is not None:
                    walk(nested)

            if value.get("type") == "text" and isinstance(value.get("text"), str):
                append(value["text"])

        walk(payload)
        return "\n".join(text_chunks).strip()

    def _render_markdown_docx(self, *, title: str, markdown: str) -> bytes:
        html = markdown_lib.markdown(
            markdown or "",
            extensions=[
                "tables",
                "fenced_code",
                "sane_lists",
                "nl2br",
            ],
        )
        soup = BeautifulSoup(html, "html.parser")
        document = Document()
        document.add_heading(title, level=0)
        self._configure_document_styles(document)
        for node in soup.contents:
            self._append_block(document, node)
        if not soup.contents:
            document.add_paragraph("")
        buffer = io.BytesIO()
        document.save(buffer)
        return buffer.getvalue()

    def _configure_document_styles(self, document: Document) -> None:
        styles = document.styles
        if "Normal" in styles:
            styles["Normal"].font.size = Pt(11)
        if "Heading 1" in styles:
            styles["Heading 1"].font.size = Pt(18)
        if "Heading 2" in styles:
            styles["Heading 2"].font.size = Pt(15)
        if "Heading 3" in styles:
            styles["Heading 3"].font.size = Pt(13)

    def _append_block(self, document: Document, node: Tag | NavigableString) -> None:
        if isinstance(node, NavigableString):
            text = str(node).strip()
            if text:
                document.add_paragraph(text)
            return
        if not isinstance(node, Tag):
            return
        if node.name in {"h1", "h2", "h3", "h4", "h5", "h6"}:
            level = min(int(node.name[1]), 4)
            paragraph = document.add_heading(level=level)
            self._append_inline(paragraph, node)
            return
        if node.name == "p":
            paragraph = document.add_paragraph()
            self._append_inline(paragraph, node)
            return
        if node.name in {"ul", "ol"}:
            self._append_list(document, node, ordered=node.name == "ol", level=0)
            return
        if node.name == "pre":
            paragraph = document.add_paragraph()
            run = paragraph.add_run(node.get_text("\n"))
            run.font.name = "Courier New"
            run.font.size = Pt(10)
            return
        if node.name == "blockquote":
            paragraph = document.add_paragraph(style="Intense Quote" if "Intense Quote" in document.styles else None)
            self._append_inline(paragraph, node)
            return
        if node.name == "table":
            self._append_table(document, node)
            return
        if node.name == "hr":
            document.add_paragraph("----------------------------------------")
            return
        paragraph = document.add_paragraph()
        self._append_inline(paragraph, node)

    def _append_list(self, document: Document, list_tag: Tag, *, ordered: bool, level: int) -> None:
        style = "List Number" if ordered else "List Bullet"
        nested_style = "List Number 2" if ordered else "List Bullet 2"
        for item in list_tag.find_all("li", recursive=False):
            paragraph = document.add_paragraph(style=nested_style if level > 0 and nested_style in document.styles else style)
            for child in item.contents:
                if isinstance(child, Tag) and child.name in {"ul", "ol"}:
                    continue
                self._append_inline(paragraph, child)
            for child in item.find_all(["ul", "ol"], recursive=False):
                self._append_list(document, child, ordered=child.name == "ol", level=level + 1)

    def _append_table(self, document: Document, table_tag: Tag) -> None:
        rows = table_tag.find_all("tr")
        if not rows:
            return
        max_cols = max(len(row.find_all(["th", "td"], recursive=False)) for row in rows)
        table = document.add_table(rows=len(rows), cols=max_cols)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.style = "Table Grid"
        for row_index, row in enumerate(rows):
            cells = row.find_all(["th", "td"], recursive=False)
            for cell_index, cell_tag in enumerate(cells):
                cell = table.cell(row_index, cell_index)
                paragraph = cell.paragraphs[0]
                self._clear_paragraph(paragraph)
                self._append_inline(paragraph, cell_tag)
                if cell_tag.name == "th":
                    for run in paragraph.runs:
                        run.bold = True

    def _append_inline(self, paragraph, node: Tag | NavigableString, *, bold: bool = False, italic: bool = False, code: bool = False) -> None:
        if isinstance(node, NavigableString):
            text = str(node)
            if text:
                run = paragraph.add_run(text)
                run.bold = bold
                run.italic = italic
                if code:
                    run.font.name = "Courier New"
                    run.font.size = Pt(10)
            return
        if not isinstance(node, Tag):
            return
        name = node.name.lower()
        if name == "br":
            paragraph.add_run("\n")
            return
        if name in {"strong", "b"}:
            for child in node.contents:
                self._append_inline(paragraph, child, bold=True or bold, italic=italic, code=code)
            return
        if name in {"em", "i"}:
            for child in node.contents:
                self._append_inline(paragraph, child, bold=bold, italic=True or italic, code=code)
            return
        if name == "code":
            for child in node.contents:
                self._append_inline(paragraph, child, bold=bold, italic=italic, code=True)
            return
        if name == "a":
            self._add_hyperlink(paragraph, node.get("href", ""), node.get_text(strip=False), bold=bold, italic=italic, code=code)
            return
        for child in node.contents:
            self._append_inline(paragraph, child, bold=bold, italic=italic, code=code)

    def _add_hyperlink(self, paragraph, url: str, text: str, *, bold: bool = False, italic: bool = False, code: bool = False) -> None:
        if not url:
            run = paragraph.add_run(text)
            run.bold = bold
            run.italic = italic
            return
        part = paragraph.part
        rel_id = part.relate_to(url, RT.HYPERLINK, is_external=True)
        hyperlink = OxmlElement("w:hyperlink")
        hyperlink.set(qn("r:id"), rel_id)
        run_element = OxmlElement("w:r")
        run_pr = OxmlElement("w:rPr")
        color = OxmlElement("w:color")
        color.set(qn("w:val"), "0563C1")
        underline = OxmlElement("w:u")
        underline.set(qn("w:val"), "single")
        run_pr.append(color)
        run_pr.append(underline)
        if bold:
            run_pr.append(OxmlElement("w:b"))
        if italic:
            run_pr.append(OxmlElement("w:i"))
        if code:
            fonts = OxmlElement("w:rFonts")
            fonts.set(qn("w:ascii"), "Courier New")
            fonts.set(qn("w:hAnsi"), "Courier New")
            run_pr.append(fonts)
        run_element.append(run_pr)
        text_element = OxmlElement("w:t")
        if text.startswith(" ") or text.endswith(" "):
            text_element.set(qn("xml:space"), "preserve")
        text_element.text = text
        run_element.append(text_element)
        hyperlink.append(run_element)
        paragraph._p.append(hyperlink)

    def _clear_paragraph(self, paragraph) -> None:
        element = paragraph._element
        for child in list(element):
            element.remove(child)

    def _build_document_xml(self, *, title: str, body_paragraphs: list[str]) -> str:
        title_paragraph, _ = self._docx_paragraph(title, style="Title")
        paragraphs = [title_paragraph]
        paragraphs.extend(body_paragraphs)
        body = "".join(paragraphs) + '<w:sectPr><w:pgSz w:w="11906" w:h="16838"/><w:pgMar w:top="1440" w:right="1440" w:bottom="1440" w:left="1440"/></w:sectPr>'
        return (
            '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
            '<w:document xmlns:wpc="http://schemas.microsoft.com/office/word/2010/wordprocessingCanvas" '
            'xmlns:mc="http://schemas.openxmlformats.org/markup-compatibility/2006" '
            'xmlns:o="urn:schemas-microsoft-com:office:office" '
            'xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships" '
            'xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math" '
            'xmlns:v="urn:schemas-microsoft-com:vml" '
            'xmlns:wp14="http://schemas.microsoft.com/office/word/2010/wordprocessingDrawing" '
            'xmlns:wp="http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing" '
            'xmlns:w10="urn:schemas-microsoft-com:office:word" '
            'xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" '
            'xmlns:w14="http://schemas.microsoft.com/office/word/2010/wordml" '
            'xmlns:wpg="http://schemas.microsoft.com/office/word/2010/wordprocessingGroup" '
            'xmlns:wpi="http://schemas.microsoft.com/office/word/2010/wordprocessingInk" '
            'xmlns:wne="http://schemas.microsoft.com/office/word/2006/wordml" '
            'xmlns:wps="http://schemas.microsoft.com/office/word/2010/wordprocessingShape" '
            'mc:Ignorable="w14 wp14"><w:body>'
            f"{body}</w:body></w:document>"
        )

    def _markdown_to_docx_paragraphs(self, markdown: str) -> tuple[list[str], list[_DocxHyperlink]]:
        lines = (markdown or "").splitlines()
        paragraphs: list[str] = []
        hyperlinks: list[_DocxHyperlink] = []
        in_code = False
        code_lines: list[str] = []
        for raw_line in lines:
            line = raw_line.rstrip()
            stripped = line.strip()
            if stripped.startswith("```"):
                if in_code:
                    paragraph, _ = self._docx_paragraph("\n".join(code_lines), style="CodeBlock")
                    paragraphs.append(paragraph)
                    code_lines = []
                in_code = not in_code
                continue
            if in_code:
                code_lines.append(line)
                continue
            if not stripped:
                paragraph, _ = self._docx_paragraph("", style="Normal")
                paragraphs.append(paragraph)
                continue
            if stripped.startswith("# "):
                paragraph, _ = self._docx_paragraph(stripped[2:].strip(), style="Heading1")
                paragraphs.append(paragraph)
                continue
            if stripped.startswith("## "):
                paragraph, _ = self._docx_paragraph(stripped[3:].strip(), style="Heading2")
                paragraphs.append(paragraph)
                continue
            if stripped.startswith("### "):
                paragraph, _ = self._docx_paragraph(stripped[4:].strip(), style="Heading3")
                paragraphs.append(paragraph)
                continue
            if re.match(r"^[-*]\s+", stripped):
                paragraph, paragraph_links = self._docx_paragraph(
                    re.sub(r"^[-*]\s+", "", stripped),
                    style="ListParagraph",
                    numbering_id=1,
                    hyperlinks=hyperlinks,
                )
                paragraphs.append(paragraph)
                hyperlinks.extend(paragraph_links)
                continue
            if re.match(r"^\d+\.\s+", stripped):
                paragraph, paragraph_links = self._docx_paragraph(
                    re.sub(r"^\d+\.\s+", "", stripped),
                    style="ListParagraph",
                    numbering_id=2,
                    hyperlinks=hyperlinks,
                )
                paragraphs.append(paragraph)
                hyperlinks.extend(paragraph_links)
                continue
            paragraph, paragraph_links = self._docx_paragraph(stripped, style="Normal", hyperlinks=hyperlinks)
            paragraphs.append(paragraph)
            hyperlinks.extend(paragraph_links)
        if in_code and code_lines:
            paragraph, _ = self._docx_paragraph("\n".join(code_lines), style="CodeBlock")
            paragraphs.append(paragraph)
        return paragraphs, hyperlinks

    def _docx_paragraph(
        self,
        text: str,
        *,
        style: str = "Normal",
        numbering_id: int | None = None,
        hyperlinks: list[_DocxHyperlink] | None = None,
    ) -> tuple[str, list[_DocxHyperlink]]:
        runs, paragraph_links = self._inline_runs_xml(text, existing_links=hyperlinks or [])
        paragraph_pr = [f"<w:pStyle w:val=\"{style}\"/>"]
        if numbering_id is not None:
            paragraph_pr.append(
                f"<w:numPr><w:ilvl w:val=\"0\"/><w:numId w:val=\"{numbering_id}\"/></w:numPr>"
            )
        return (
            "<w:p>"
            f"<w:pPr>{''.join(paragraph_pr)}</w:pPr>"
            f"{runs}"
            "</w:p>",
            paragraph_links,
        )

    def _inline_runs_xml(self, text: str, *, existing_links: list[_DocxHyperlink]) -> tuple[str, list[_DocxHyperlink]]:
        segments = self._parse_inline_segments(text)
        runs: list[str] = []
        new_links: list[_DocxHyperlink] = []
        for segment in segments:
            if segment["type"] == "link":
                rel_id = f"rId{len(existing_links) + len(new_links) + 100}"
                link = _DocxHyperlink(rel_id=rel_id, url=segment["url"])
                new_links.append(link)
                link_runs = self._text_runs_xml(segment["text"], hyperlink=True)
                runs.append(f'<w:hyperlink r:id="{rel_id}">{link_runs}</w:hyperlink>')
                continue
            segment_text = segment["text"]
            is_bold = segment["bold"]
            is_code = segment["code"]
            properties: list[str] = []
            if is_bold:
                properties.append("<w:b/>")
            if is_code:
                properties.extend(
                    [
                        '<w:rFonts w:ascii="Courier New" w:hAnsi="Courier New" w:eastAsia="Courier New"/>',
                        '<w:shd w:val="clear" w:color="auto" w:fill="EDEDED"/>',
                    ]
                )
            text_xml = escape(segment_text)
            if segment_text.startswith(" ") or segment_text.endswith(" ") or "\n" in segment_text:
                runs.append(f"<w:r><w:rPr>{''.join(properties)}</w:rPr><w:t xml:space=\"preserve\">{text_xml}</w:t></w:r>")
            else:
                runs.append(f"<w:r><w:rPr>{''.join(properties)}</w:rPr><w:t>{text_xml}</w:t></w:r>")
        return "".join(runs) or "<w:r><w:t></w:t></w:r>", new_links

    def _text_runs_xml(self, text: str, *, hyperlink: bool = False) -> str:
        segments = self._parse_text_segments(text)
        runs: list[str] = []
        for segment_text, is_bold, is_code in segments:
            properties: list[str] = []
            if is_bold:
                properties.append("<w:b/>")
            if hyperlink:
                properties.extend(["<w:u w:val=\"single\"/>", '<w:color w:val="0563C1"/>'])
            if is_code:
                properties.extend(
                    [
                        '<w:rFonts w:ascii="Courier New" w:hAnsi="Courier New" w:eastAsia="Courier New"/>',
                        '<w:shd w:val="clear" w:color="auto" w:fill="EDEDED"/>',
                    ]
                )
            text_xml = escape(segment_text)
            if segment_text.startswith(" ") or segment_text.endswith(" ") or "\n" in segment_text:
                runs.append(f"<w:r><w:rPr>{''.join(properties)}</w:rPr><w:t xml:space=\"preserve\">{text_xml}</w:t></w:r>")
            else:
                runs.append(f"<w:r><w:rPr>{''.join(properties)}</w:rPr><w:t>{text_xml}</w:t></w:r>")
        return "".join(runs) or "<w:r><w:t></w:t></w:r>"

    def _parse_inline_segments(self, text: str) -> list[dict[str, str | bool]]:
        segments: list[dict[str, str | bool]] = []
        index = 0
        link_pattern = re.compile(r"\[([^\]]+)\]\((https?://[^)\s]+)\)")
        for match in link_pattern.finditer(text or ""):
            if match.start() > index:
                segments.extend(self._plain_text_segment_dicts(text[index:match.start()]))
            segments.append({"type": "link", "text": match.group(1), "url": match.group(2)})
            index = match.end()
        if index < len(text or ""):
            segments.extend(self._plain_text_segment_dicts(text[index:]))
        return segments or [{"type": "text", "text": "", "bold": False, "code": False}]

    def _plain_text_segment_dicts(self, text: str) -> list[dict[str, str | bool]]:
        return [
            {"type": "text", "text": segment_text, "bold": is_bold, "code": is_code}
            for segment_text, is_bold, is_code in self._parse_text_segments(text)
        ]

    def _parse_text_segments(self, text: str) -> list[tuple[str, bool, bool]]:
        if not text:
            return [("", False, False)]
        segments: list[tuple[str, bool, bool]] = []
        index = 0
        while index < len(text):
            if text.startswith("**", index):
                end = text.find("**", index + 2)
                if end != -1:
                    segments.append((text[index + 2 : end], True, False))
                    index = end + 2
                    continue
            if text.startswith("`", index):
                end = text.find("`", index + 1)
                if end != -1:
                    segments.append((text[index + 1 : end], False, True))
                    index = end + 1
                    continue
            next_markers = [pos for pos in (text.find("**", index), text.find("`", index)) if pos != -1]
            next_index = min(next_markers) if next_markers else len(text)
            segments.append((text[index:next_index], False, False))
            index = next_index
        return [segment for segment in segments if segment[0] or segment[1] or segment[2]]

    def _content_types_xml(self) -> str:
        return (
            '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
            '<Types xmlns="http://schemas.openxmlformats.org/package/2006/content-types">'
            '<Default Extension="rels" ContentType="application/vnd.openxmlformats-package.relationships+xml"/>'
            '<Default Extension="xml" ContentType="application/xml"/>'
            '<Override PartName="/word/document.xml" ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.document.main+xml"/>'
            '<Override PartName="/word/styles.xml" ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.styles+xml"/>'
            '<Override PartName="/word/numbering.xml" ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.numbering+xml"/>'
            '<Override PartName="/docProps/core.xml" ContentType="application/vnd.openxmlformats-package.core-properties+xml"/>'
            '<Override PartName="/docProps/app.xml" ContentType="application/vnd.openxmlformats-officedocument.extended-properties+xml"/>'
            "</Types>"
        )

    def _root_relationships_xml(self) -> str:
        return (
            '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
            '<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">'
            '<Relationship Id="rId1" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/officeDocument" Target="word/document.xml"/>'
            '<Relationship Id="rId2" Type="http://schemas.openxmlformats.org/package/2006/relationships/metadata/core-properties" Target="docProps/core.xml"/>'
            '<Relationship Id="rId3" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/extended-properties" Target="docProps/app.xml"/>'
            "</Relationships>"
        )

    def _document_relationships_xml(self, hyperlinks: list[_DocxHyperlink]) -> str:
        link_xml = "".join(
            f'<Relationship Id="{link.rel_id}" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink" Target="{escape(link.url)}" TargetMode="External"/>'
            for link in hyperlinks
        )
        return (
            '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
            '<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">'
            '<Relationship Id="rId1" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/styles" Target="styles.xml"/>'
            '<Relationship Id="rId2" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/numbering" Target="numbering.xml"/>'
            f"{link_xml}"
            "</Relationships>"
        )

    def _numbering_xml(self) -> str:
        return (
            '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
            '<w:numbering xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
            '<w:abstractNum w:abstractNumId="0">'
            '<w:multiLevelType w:val="hybridMultilevel"/>'
            '<w:lvl w:ilvl="0"><w:start w:val="1"/><w:numFmt w:val="bullet"/>'
            '<w:lvlText w:val="•"/><w:lvlJc w:val="left"/>'
            '<w:pPr><w:ind w:left="720" w:hanging="360"/></w:pPr></w:lvl>'
            '</w:abstractNum>'
            '<w:abstractNum w:abstractNumId="1">'
            '<w:multiLevelType w:val="hybridMultilevel"/>'
            '<w:lvl w:ilvl="0"><w:start w:val="1"/><w:numFmt w:val="decimal"/>'
            '<w:lvlText w:val="%1."/><w:lvlJc w:val="left"/>'
            '<w:pPr><w:ind w:left="720" w:hanging="360"/></w:pPr></w:lvl>'
            '</w:abstractNum>'
            '<w:num w:numId="1"><w:abstractNumId w:val="0"/></w:num>'
            '<w:num w:numId="2"><w:abstractNumId w:val="1"/></w:num>'
            "</w:numbering>"
        )

    def _styles_xml(self) -> str:
        return (
            '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
            '<w:styles xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
            '<w:style w:type="paragraph" w:default="1" w:styleId="Normal"><w:name w:val="Normal"/></w:style>'
            '<w:style w:type="paragraph" w:styleId="Title"><w:name w:val="Title"/><w:rPr><w:b/><w:sz w:val="32"/></w:rPr></w:style>'
            '<w:style w:type="paragraph" w:styleId="Heading1"><w:name w:val="heading 1"/><w:rPr><w:b/><w:sz w:val="30"/></w:rPr></w:style>'
            '<w:style w:type="paragraph" w:styleId="Heading2"><w:name w:val="heading 2"/><w:rPr><w:b/><w:sz w:val="26"/></w:rPr></w:style>'
            '<w:style w:type="paragraph" w:styleId="Heading3"><w:name w:val="heading 3"/><w:rPr><w:b/><w:sz w:val="24"/></w:rPr></w:style>'
            '<w:style w:type="paragraph" w:styleId="ListParagraph"><w:name w:val="List Paragraph"/></w:style>'
            '<w:style w:type="character" w:styleId="Hyperlink"><w:name w:val="Hyperlink"/><w:rPr><w:color w:val="0563C1"/><w:u w:val="single"/></w:rPr></w:style>'
            '<w:style w:type="paragraph" w:styleId="CodeBlock"><w:name w:val="Code Block"/><w:rPr>'
            '<w:rFonts w:ascii="Courier New" w:hAnsi="Courier New" w:eastAsia="Courier New"/><w:sz w:val="20"/>'
            '</w:rPr></w:style>'
            "</w:styles>"
        )

    def _build_core_xml(self, *, title: str) -> str:
        safe_title = escape(title)
        return (
            '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
            '<cp:coreProperties xmlns:cp="http://schemas.openxmlformats.org/package/2006/metadata/core-properties" '
            'xmlns:dc="http://purl.org/dc/elements/1.1/" '
            'xmlns:dcterms="http://purl.org/dc/terms/" '
            'xmlns:dcmitype="http://purl.org/dc/dcmitype/" '
            'xmlns:xsi="http://www.w3.org/2001/XMLSchema-instance">'
            f"<dc:title>{safe_title}</dc:title>"
            "<dc:creator>Codex</dc:creator>"
            "<cp:lastModifiedBy>Codex</cp:lastModifiedBy>"
            "</cp:coreProperties>"
        )

    def _app_xml(self) -> str:
        return (
            '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
            '<Properties xmlns="http://schemas.openxmlformats.org/officeDocument/2006/extended-properties" '
            'xmlns:vt="http://schemas.openxmlformats.org/officeDocument/2006/docPropsVTypes">'
            "<Application>Codex</Application>"
            "</Properties>"
        )
